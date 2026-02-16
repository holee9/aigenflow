"""
Output formatters for document export.

Provides formatters for Markdown, DOCX, and PDF output formats.
"""

import io
from abc import ABC, abstractmethod
from enum import StrEnum
from typing import Any

from core.logger import get_logger

logger = get_logger(__name__)


class OutputFormat(StrEnum):
    """Supported output formats."""

    MARKDOWN = "md"
    DOCX = "docx"
    PDF = "pdf"
    ALL = "all"


class OutputFormatter(ABC):
    """
    Abstract interface for document formatters.

    All formatters must implement format_document and get_format methods.
    """

    @abstractmethod
    def format_document(
        self,
        content: str,
        metadata: dict[str, Any] | None = None,
    ) -> str | bytes:
        """
        Format content as a document.

        Args:
            content: Document content (typically Markdown)
            metadata: Optional metadata (title, author, date, etc.)

        Returns:
            Formatted document as string (for text formats) or bytes (for binary formats)
        """
        raise NotImplementedError

    @abstractmethod
    def get_format(self) -> OutputFormat:
        """
        Get the output format for this formatter.

        Returns:
            OutputFormat enum value
        """
        raise NotImplementedError


class MarkdownFormatter(OutputFormatter):
    """
    Markdown output formatter.

    Formats documents as Markdown text with optional frontmatter.
    """

    def format_document(
        self,
        content: str,
        metadata: dict[str, Any] | None = None,
    ) -> str:
        """
        Format content as markdown document.

        Args:
            content: Document content in Markdown
            metadata: Optional metadata (title, author, date, etc.)

        Returns:
            Markdown string with optional YAML frontmatter
        """
        if not content:
            return ""

        # Add frontmatter if metadata provided
        if metadata:
            frontmatter = self._create_frontmatter(metadata)
            return f"{frontmatter}\n\n{content}"

        return content

    def get_format(self) -> OutputFormat:
        """Get the output format."""
        return OutputFormat.MARKDOWN

    def _create_frontmatter(self, metadata: dict[str, Any]) -> str:
        """Create YAML frontmatter from metadata."""
        lines = ["---"]
        for key, value in metadata.items():
            if key == "tags" and isinstance(value, list):
                lines.append(f"{key}: [{', '.join(value)}]")
            elif isinstance(value, str):
                lines.append(f'{key}: "{value}"')
            else:
                lines.append(f"{key}: {value}")
        lines.append("---")
        return "\n".join(lines)


class DocxFormatter(OutputFormatter):
    """
    DOCX output formatter using python-docx.

    Formats documents as Microsoft Word .docx files.
    """

    def format_document(
        self,
        content: str,
        metadata: dict[str, Any] | None = None,
    ) -> bytes:
        """
        Format content as DOCX document.

        Args:
            content: Document content in Markdown
            metadata: Optional metadata (title, author, date, etc.)

        Returns:
            DOCX file content as bytes
        """
        try:
            from docx import Document
        except ImportError:
            logger.warning("python-docx not installed, returning empty DOCX")
            return b""

        doc = Document()

        # Add title from metadata if provided
        title = metadata.get("title", "Document") if metadata else "Document"
        doc.add_heading(title, 0)

        # Add metadata if provided
        if metadata:
            if "author" in metadata:
                doc.add_paragraph(f"Author: {metadata['author']}")
            if "date" in metadata:
                doc.add_paragraph(f"Date: {metadata['date']}")
            doc.add_paragraph()  # Empty line separator

        # Parse and add content
        self._add_markdown_content(doc, content)

        # Write to bytes buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        return buffer.getvalue()

    def get_format(self) -> OutputFormat:
        """Get the output format."""
        return OutputFormat.DOCX

    def _add_markdown_content(self, doc: Any, content: str) -> None:
        """
        Parse Markdown and add to document.

        Args:
            doc: python-docx Document object
            content: Markdown content string
        """

        lines = content.split("\n")
        in_list = False

        for line in lines:
            line = line.rstrip()

            # Handle headings
            if line.startswith("#"):
                in_list = False
                level = min(len(line) - len(line.lstrip("#")), 6)
                text = line.lstrip("#").strip()
                doc.add_heading(text, level=level)

            # Handle horizontal rules
            elif line.strip() == "---":
                in_list = False
                pass  # DOCX doesn't have horizontal rules

            # Handle lists
            elif line.strip().startswith(("- ", "* ", "+ ")):
                if not in_list:
                    in_list = True
                text = line.strip()[2:]
                doc.add_paragraph(text, style="List Bullet")

            # Handle numbered lists
            elif line.strip().startswith(tuple(f"{i}. " for i in range(1, 10))):
                if not in_list:
                    in_list = True
                text = line.strip().split(".", 1)[1].strip()
                doc.add_paragraph(text, style="List Number")

            # Handle empty lines
            elif not line.strip():
                in_list = False
                doc.add_paragraph()

            # Regular paragraph
            else:
                in_list = False
                # Parse inline formatting (bold, italic)
                paragraph = doc.add_paragraph()
                self._add_formatted_text(paragraph, line)

    def _add_formatted_text(self, paragraph: Any, text: str) -> None:
        """
        Add text with inline formatting to a paragraph.

        Args:
            paragraph: python-docx Paragraph object
            text: Text with potential inline formatting
        """
        from docx.shared import Pt

        # Simple parsing for **bold** and *italic*
        parts = []
        current = ""
        i = 0

        while i < len(text):
            # Check for bold **text**
            if text[i:i + 2] == "**":
                if current:
                    parts.append(("text", current))
                current = ""
                i += 2
                end = text.find("**", i)
                if end != -1:
                    bold_text = text[i:end]
                    parts.append(("bold", bold_text))
                    i = end + 2
                else:
                    current = "**" + text[i:]
                    i = len(text)

            # Check for italic *text*
            elif text[i] == "*":
                if current:
                    parts.append(("text", current))
                current = ""
                i += 1
                end = text.find("*", i)
                if end != -1:
                    italic_text = text[i:end]
                    parts.append(("italic", italic_text))
                    i = end + 1
                else:
                    current = "*" + text[i:]
                    i = len(text)

            else:
                current += text[i]
                i += 1

        if current:
            parts.append(("text", current))

        # Add runs to paragraph
        for style, content in parts:
            if content:
                run = paragraph.add_run(content)
                run.font.size = Pt(11)
                if style == "bold":
                    run.bold = True
                elif style == "italic":
                    run.italic = True


class PdfFormatter(OutputFormatter):
    """
    PDF output formatter using reportlab.

    Formats documents as PDF files.
    """

    def format_document(
        self,
        content: str,
        metadata: dict[str, Any] | None = None,
    ) -> bytes:
        """
        Format content as PDF document.

        Args:
            content: Document content in Markdown
            metadata: Optional metadata (title, author, date, etc.)

        Returns:
            PDF file content as bytes
        """
        try:
            from reportlab.lib.enums import TA_CENTER
            from reportlab.lib.pagesizes import A4
            from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
            from reportlab.lib.units import inch
            from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
        except ImportError:
            logger.warning("reportlab not installed, returning empty PDF")
            return b""

        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4)

        # Container for PDF elements
        story = []
        styles = getSampleStyleSheet()

        # Create custom styles
        title_style = ParagraphStyle(
            "CustomTitle",
            parent=styles["Heading1"],
            fontSize=24,
            textColor="#333333",
            alignment=TA_CENTER,
            spaceAfter=30,
        )

        heading_style = ParagraphStyle(
            "CustomHeading",
            parent=styles["Heading2"],
            fontSize=16,
            textColor="#333333",
            spaceAfter=12,
        )

        body_style = ParagraphStyle(
            "CustomBody",
            parent=styles["BodyText"],
            fontSize=11,
            leading=14,
            spaceAfter=12,
        )

        # Add title
        title = metadata.get("title", "Document") if metadata else "Document"
        story.append(Paragraph(title, title_style))
        story.append(Spacer(1, 0.2 * inch))

        # Add metadata
        if metadata:
            if "author" in metadata:
                story.append(Paragraph(f"<b>Author:</b> {metadata['author']}", body_style))
            if "date" in metadata:
                story.append(Paragraph(f"<b>Date:</b> {metadata['date']}", body_style))
            story.append(Spacer(1, 0.2 * inch))

        # Parse and add content
        self._add_markdown_to_story(story, content, styles, body_style, heading_style)

        # Build PDF
        doc.build(story)
        return buffer.getvalue()

    def get_format(self) -> OutputFormat:
        """Get the output format."""
        return OutputFormat.PDF

    def _add_markdown_to_story(
        self,
        story: list,
        content: str,
        styles: Any,
        body_style: Any,
        heading_style: Any,
    ) -> None:
        """Parse Markdown and add to PDF story."""
        from reportlab.lib.units import inch
        from reportlab.platypus import Paragraph, Spacer

        lines = content.split("\n")
        in_list = False
        list_items = []

        for line in lines:
            line = line.rstrip()

            # Handle headings
            if line.startswith("#"):
                if in_list and list_items:
                    self._add_list_to_story(story, list_items, body_style)
                    list_items = []
                    in_list = False

                level = min(len(line) - len(line.lstrip("#")), 6)
                text = self._clean_markdown(line.lstrip("#").strip())
                if level == 1:
                    story.append(Paragraph(text, heading_style))
                else:
                    style = getattr(styles, f"Heading{level}", body_style)
                    story.append(Paragraph(text, style))

            # Handle list items
            elif line.strip().startswith(("- ", "* ", "+ ")):
                in_list = True
                text = self._clean_markdown(line.strip()[2:])
                list_items.append(text)

            # Handle numbered lists
            elif line.strip().split(".")[0].isdigit() and len(line.strip()) > 2:
                parts = line.strip().split(".", 1)
                if len(parts) == 2:
                    in_list = True
                    text = self._clean_markdown(parts[1].strip())
                    list_items.append(text)

            # Handle empty lines or end of list
            elif not line.strip():
                if in_list and list_items:
                    self._add_list_to_story(story, list_items, body_style)
                    list_items = []
                    in_list = False
                story.append(Spacer(1, 0.1 * inch))

            # Regular paragraph
            else:
                if in_list and list_items:
                    self._add_list_to_story(story, list_items, body_style)
                    list_items = []
                    in_list = False

                text = self._clean_markdown(line)
                if text:
                    story.append(Paragraph(text, body_style))

        # Add remaining list items
        if in_list and list_items:
            self._add_list_to_story(story, list_items, body_style)

    def _clean_markdown(self, text: str) -> str:
        """Clean Markdown syntax for PDF."""
        # Remove bold/italic markers
        text = text.replace("**", "").replace("__", "")
        text = text.replace("*", "").replace("_", "")
        # Escape special characters
        text = text.replace("<", "&lt;").replace(">", "&gt;")
        return text

    def _add_list_to_story(self, story: list, items: list[str], style: Any) -> None:
        """Add list items to story."""
        from reportlab.lib.units import inch
        from reportlab.platypus import Paragraph, Spacer

        for item in items:
            story.append(Paragraph(f"&#8226; {item}", style))
        story.append(Spacer(1, 0.1 * inch))


def get_formatter(format_type: OutputFormat | str) -> OutputFormatter:
    """
    Get formatter instance for specified format.

    Args:
        format_type: OutputFormat enum or string value

    Returns:
        OutputFormatter instance

    Raises:
        ValueError: If format type is not supported
    """
    if isinstance(format_type, str):
        try:
            format_type = OutputFormat(format_type)
        except ValueError:
            raise ValueError(f"Unknown output format: {format_type}")

    if format_type == OutputFormat.MARKDOWN:
        return MarkdownFormatter()
    elif format_type == OutputFormat.DOCX:
        return DocxFormatter()
    elif format_type == OutputFormat.PDF:
        return PdfFormatter()
    else:
        raise ValueError(f"Unsupported format type: {format_type}")
